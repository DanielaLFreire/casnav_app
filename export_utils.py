# -*- coding: utf-8 -*-
"""
CASNAV DMarSup - Exportacao de Relatorios (MD + DOCX)
Versao 3.0 - Sem PDF - Compativel Python 3.8+
Dependencia: pip install python-docx
"""

import io
import re
from datetime import datetime

__version__ = "3.0"


def _parse_md_lines(md_text):
    elements = []
    for line in md_text.split("\n"):
        s = line.strip()
        if not s:
            continue
        if s.startswith("####"):
            elements.append(("h4", s.lstrip("#").strip()))
        elif s.startswith("###"):
            elements.append(("h3", s.lstrip("#").strip()))
        elif s.startswith("##"):
            elements.append(("h2", s.lstrip("#").strip()))
        elif s.startswith("#"):
            elements.append(("h1", s.lstrip("#").strip()))
        elif s in ("---", "***", "___"):
            elements.append(("hr", ""))
        elif s.startswith("|") and s.endswith("|"):
            cells = [c.strip() for c in s.split("|")[1:-1]]
            if all(re.match(r'^[-:]+$', c) for c in cells):
                continue
            elements.append(("table_row", cells))
        elif s.startswith("- "):
            elements.append(("bullet", s[2:].strip()))
        elif s.startswith("**") and s.endswith("**"):
            elements.append(("bold_para", s.strip("*").strip()))
        else:
            elements.append(("para", s))
    return elements


def _clean_bold(text):
    parts = []
    pattern = re.compile(r'\*\*(.+?)\*\*')
    last = 0
    for m in pattern.finditer(text):
        if m.start() > last:
            parts.append((text[last:m.start()], False))
        parts.append((m.group(1), True))
        last = m.end()
    if last < len(text):
        parts.append((text[last:], False))
    if not parts:
        return [(text, False)]
    return parts


def gerar_docx(md_text, titulo="Relatorio CASNAV DMarSup"):
    from docx import Document
    from docx.shared import Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT

    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Arial'
    style.font.size = Pt(11)

    for sec in doc.sections:
        sec.top_margin = Cm(2.5)
        sec.bottom_margin = Cm(2.5)
        sec.left_margin = Cm(2.5)
        sec.right_margin = Cm(2.5)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(titulo)
    r.bold = True
    r.font.size = Pt(14)
    r.font.color.rgb = RGBColor(0x1A, 0x3C, 0x6E)

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run("Projeto Sistemas Maritimos Nao Tripulados")
    r2.font.size = Pt(10)
    r2.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r3 = p3.add_run("Gerado em: %s" % datetime.now().strftime('%d/%m/%Y %H:%M'))
    r3.font.size = Pt(9)
    r3.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    doc.add_paragraph()

    elements = _parse_md_lines(md_text)
    i = 0
    while i < len(elements):
        etype, content = elements[i]

        if etype in ("h1", "h2", "h3", "h4"):
            level = int(etype[1])
            doc.add_heading(content, level=level)
        elif etype == "hr":
            p = doc.add_paragraph()
            r = p.add_run("_" * 60)
            r.font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)
        elif etype == "bullet":
            p = doc.add_paragraph(style='List Bullet')
            for txt, bold in _clean_bold(content):
                r = p.add_run(txt)
                r.bold = bold
                r.font.size = Pt(11)
        elif etype == "bold_para":
            p = doc.add_paragraph()
            r = p.add_run(content)
            r.bold = True
            r.font.size = Pt(11)
        elif etype == "table_row":
            rows = [content]
            j = i + 1
            while j < len(elements) and elements[j][0] == "table_row":
                rows.append(elements[j][1])
                j += 1
            ncols = max(len(row) for row in rows)
            table = doc.add_table(rows=len(rows), cols=ncols)
            table.style = 'Light Grid Accent 1'
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            for ri, row_data in enumerate(rows):
                for ci, cell_text in enumerate(row_data):
                    if ci < ncols:
                        cell = table.cell(ri, ci)
                        cell.text = ""
                        p = cell.paragraphs[0]
                        for txt, bold in _clean_bold(cell_text):
                            r = p.add_run(txt)
                            r.bold = bold or (ri == 0)
                            r.font.size = Pt(10)
                            r.font.name = 'Arial'
            doc.add_paragraph()
            i = j
            continue
        else:
            p = doc.add_paragraph()
            for txt, bold in _clean_bold(content):
                r = p.add_run(txt)
                r.bold = bold
                r.font.size = Pt(11)
        i += 1

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.getvalue()


def render_botoes_download(md_text, filename_base, titulo="Relatorio CASNAV DMarSup"):
    """Renderiza 2 botoes de download (MD + DOCX) no Streamlit."""
    import streamlit as st

    col1, col2 = st.columns(2)

    with col1:
        st.download_button(
            label="📥 Markdown (.md)",
            data=md_text,
            file_name="%s.md" % filename_base,
            mime="text/markdown",
            use_container_width=True
        )

    with col2:
        docx_bytes = gerar_docx(md_text, titulo)
        st.download_button(
            label="📥 Word (.docx)",
            data=docx_bytes,
            file_name="%s.docx" % filename_base,
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            use_container_width=True
        )
